"""
Document parsing service for extracting text from PDF and DOCX files.
"""
from pathlib import Path
from typing import Optional
import re
from pypdf import PdfReader
from docx import Document
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentParser:
    """Service for parsing documents and extracting text."""
    
    @staticmethod
    def parse_pdf(file_path: Path) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
            
        Raises:
            Exception: If parsing fails
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(
                        f"Error extracting text from page {page_num}: {e}",
                        extra={"extra_fields": {"file_path": str(file_path), "page": page_num}}
                    )
            
            full_text = "\n".join(text_parts)
            normalized_text = DocumentParser._normalize_text(full_text)
            
            logger.info(
                f"Parsed PDF: {len(reader.pages)} pages, {len(normalized_text)} characters",
                extra={"extra_fields": {"file_path": str(file_path), "pages": len(reader.pages)}}
            )
            
            return normalized_text
            
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}", extra={"extra_fields": {"file_path": str(file_path)}})
            raise
    
    @staticmethod
    def parse_docx(file_path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
            
        Raises:
            Exception: If parsing fails
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            normalized_text = DocumentParser._normalize_text(full_text)
            
            logger.info(
                f"Parsed DOCX: {len(normalized_text)} characters",
                extra={"extra_fields": {"file_path": str(file_path)}}
            )
            
            return normalized_text
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}", extra={"extra_fields": {"file_path": str(file_path)}})
            raise
    
    @staticmethod
    def parse_file(file_path: Path) -> str:
        """
        Parse a file based on its extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If file type is not supported
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == ".pdf":
            return DocumentParser.parse_pdf(file_path)
        elif file_ext in [".doc", ".docx"]:
            return DocumentParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    @staticmethod
    def _normalize_text(text: str) -> str:
        """
        Normalize whitespace and clean up text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Normalized text
        """
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = text.strip()
        
        return text

